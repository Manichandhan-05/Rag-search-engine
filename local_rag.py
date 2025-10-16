import os
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
import re

class LocalKnowledgeBase:
    def __init__(self):
        self.documents = []  # Store documents as simple text
        print("✓ Local Knowledge Base initialized (No API key required)")

    def add_documents(self, file_paths):
        documents = []
        for file_path in file_paths:
            print(f"Processing file: {file_path}")
            if file_path.endswith('.pdf'):
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            elif file_path.endswith('.docx'):
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            else:
                print(f"Skipping unsupported file: {file_path}")
                continue

        # Store documents for simple retrieval
        self.documents.extend(documents)
        print(f"✓ Added {len(documents)} documents to knowledge base")

    def query(self, question):
        if not self.documents:
            return "No documents uploaded yet."
        
        # Simple keyword-based search
        question_lower = question.lower()
        question_words = set(re.findall(r'\b\w+\b', question_lower))
        
        best_match = None
        best_score = 0
        
        for doc in self.documents:
            content_lower = doc.page_content.lower()
            content_words = set(re.findall(r'\b\w+\b', content_lower))
            
            # Calculate relevance score
            common_words = question_words.intersection(content_words)
            score = len(common_words)
            
            if score > best_score:
                best_score = score
                best_match = doc
        
        if best_match is None:
            return "I couldn't find relevant information in the uploaded documents to answer your question."
        
        # Extract relevant sentences
        sentences = re.split(r'[.!?]+', best_match.page_content)
        relevant_sentences = []
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(word in sentence_lower for word in question_words if len(word) > 3):
                relevant_sentences.append(sentence.strip())
        
        if not relevant_sentences:
            # If no specific sentences match, return first few sentences
            relevant_sentences = sentences[:3]
        
        # Limit response length
        response = ". ".join(relevant_sentences[:5])
        if len(response) > 500:
            response = response[:500] + "..."
        
        return f"Based on the uploaded documents:\n\n{response}"
